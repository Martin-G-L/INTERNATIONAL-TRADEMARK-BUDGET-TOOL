import tkinter as tk
from tkinter import messagebox, ttk
import pandas as pd
import unicodedata
from tabulate import tabulate
import os
from datetime import datetime
import pytz
import re
import subprocess
import shutil
from docx import Document
import logging
import shutil

# Constants
BASE_PATH = r"C:\Users\Martin Lartigue\Documents\PROGRAM TESTING FOLDERS\exes\Project Schedule of Fees\CODE"
output_dir = os.path.join(BASE_PATH, "OUTPUT")
os.makedirs(output_dir, exist_ok=True)
logs_dir = os.path.join(output_dir, "LOGS")
os.makedirs(logs_dir, exist_ok=True)
pdf_output_dir = output_dir  # PDFs will be saved in OUTPUT folder
os.makedirs(pdf_output_dir, exist_ok=True)

FILE_PATH = os.path.join(BASE_PATH, "PLANILHAS", "SCHEDULE OF FEES.xlsx")
DOCX_TEMPLATE_PATH = os.path.join(BASE_PATH, "PLANILHAS", "Word_Template_Deposito_1_Classe.docx")
DOCX_MULTI_CLASS_TEMPLATE_PATH = os.path.join(BASE_PATH, "PLANILHAS", "Word_Template_Deposito_Multi_Classe.docx")
DOCX_PRORROGACAO_ORDINARIO_TEMPLATE_PATH = os.path.join(BASE_PATH, "PLANILHAS", "Word_Template_Prorrogacao_Ordinario_1_Classe.docx")
DOCX_PRORROGACAO_ORDINARIO_MULTI_CLASS_TEMPLATE_PATH = os.path.join(BASE_PATH, "PLANILHAS", "Word_Template_Prorrogacao_Ordinario_Multi_Classe.docx")
DOCX_PRORROGACAO_EXTRA_ORDINARIO_TEMPLATE_PATH = os.path.join(BASE_PATH, "PLANILHAS", "Word_Template_Prorrogacao_Extra_Ordinario_1_Classe.docx")
DOCX_PRORROGACAO_EXTRA_ORDINARIO_MULTI_CLASS_TEMPLATE_PATH = os.path.join(BASE_PATH, "PLANILHAS", "Word_Template_Prorrogacao_Extra_Ordinario_Multi_Classe.docx")
LOGO_PATH = os.path.join(BASE_PATH, "Logo", "logo.png")

month_names = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

honorarios_dict = {
    "Honorarios_Busca": 120,
    "Honorarios_Deposito_de_Marca_Mista_1_Classe": 410,
    "Honorarios_Deposito_de_Marca_Mista_Classe_Adicional": 200,
    "Honorarios_Deposito_de_Marca_Nominativa_1_Classe": 380,
    "Honorarios_Deposito_de_Marca_Nominativa_Classe_Adicional": 200,
    "Honorarios_Deposito_de_Marca_Figurativa_1_Classe": 380,
    "Honorarios_Deposito_de_Marca_Figurativa_Classe_Adicional": 200,
    "Honorarios_Oposicao_a_Deposito_de_Marca": 350,
    "Honorarios_Concessao_de_Marca_1_Classe": 350,
    "Honorarios_Concessao_de_Marca_Classe_Adicional": 300,
    "Honorarios_Prorrogacao_Prazo_Ordinario_1_Classe": 500,
    "Honorarios_Prorrogacao_Prazo_Extra_Ordinario_1_Classe": 700,
    "Honorarios_Prorrogacao_Prazo_Ordinario_Classe_Adicional": 300,
    "Honorarios_Prorrogacao_Prazo_Extra_Ordinario_Classe_Adicional": 600,
}

# LaTeX templates (unchanged)
latex_template_deposito = """
\\documentclass[a4paper,12pt]{article}
\\usepackage{fontspec}
\\setmainfont{Arial}
\\usepackage{polyglossia}
\\setmainlanguage{portuguese}
\\usepackage{geometry}
\\geometry{margin=2.2cm, headsep=1.5cm}
\\usepackage{graphicx}
\\usepackage{xcolor}
\\usepackage{colortbl}
\\usepackage{fancyhdr}
\\usepackage{parskip}
\\usepackage{lastpage}
\\usepackage{array}
\\usepackage{booktabs}
\\definecolor{readerred}{RGB}{207,45,75}
\\setlength{\\headheight}{70pt}
\\addtolength{\\topmargin}{-10pt}
\\pagestyle{fancy}
\\fancyhf{}
\\fancyhead[L]{\\color{readerred}\\small\\textbf{Insert Company Name} \\\\[-0.1cm] CNPJ: 12.345.567/0000-01 \\\\[-0.1cm] Insert Company Address Zip Code: 000-00 \\\\[-0.1cm] Company Phone Contact | Company Email Contact}
\\fancyhead[R]{\\raisebox{-0.5cm}{\\includegraphics[width=0.3\\textwidth]{logo.png}}}
\\fancyfoot[C]{\\thepage\\ de \\pageref{LastPage}}
\\newcolumntype{R}[1]{>{\\raggedleft\\arraybackslash}p{#1}}
\\newcolumntype{L}[1]{>{\\raggedright\\arraybackslash}p{#1}}
\\begin{document}
%s
\\end{document}
"""

latex_template_deposito_page = """
\\vspace*{-0.25cm}
\\begin{center}
    {\\LARGE \\textbf{Orçamento em %s}}\\\\
    \\vspace{0.5cm}
    {\\large Data: %s}
\\end{center}
\\section*{Informações do Cliente}
\\begin{tabular}{ll}
    \\textbf{Titular:} & %s \\\\
    \\textbf{Marca:} & %s \\\\
\\end{tabular}
\\section*{Detalhamento do Orçamento}
\\begin{tabular}{|L{10cm}|R{3cm}|R{3cm}|}
    \\hline
    \\rowcolor{readerred}
    \\color{white}\\textbf{Descrição} & \\color{white}\\textbf{Taxa do Correspondente (US\\$)} & \\color{white}\\textbf{Valor Total (US\\$)} \\\\
    \\hline
%s
    \\multicolumn{2}{|r|}{\\textbf{Total Estimado Até Registro:}} & \\textbf{%s} \\\\
    \\hline
\\end{tabular}
\\section*{Observações}
Os valores acima informados referem-se às etapas principais do depósito de marca, por marca e por classe:
\\begin{itemize}
    \\item Primeira etapa: Busca,
    \\item Segunda etapa: Depósito,
    \\item Terceira etapa (quando ocorrer): Registro, expedição do Certificado de Registro, Acompanhamento e Vigilância por 10 anos.
\\end{itemize}
Disclaimer Text Explaining Trademark Application Process.
\\newpage
"""

latex_template_prorrogacao_page = """
\\vspace*{-0.25cm}
\\begin{center}
    {\\LARGE \\textbf{Orçamento em %s}}\\\\
    \\vspace{0.5cm}
    {\\large Data: %s}
\\end{center}
\\section*{Informações do Cliente}
\\begin{tabular}{ll}
    \\textbf{Titular:} & %s \\\\
    \\textbf{Marca:} & %s \\\\
\\end{tabular}
\\section*{Detalhamento do Orçamento}
\\begin{tabular}{|L{10cm}|R{3cm}|R{3cm}|}
    \\hline
    \\rowcolor{readerred}
    \\color{white}\\textbf{Descrição} & \\color{white}\\textbf{Taxa do Correspondente (US\\$)} & \\color{white}\\textbf{Valor Total (US\\$)} \\\\
    \\hline
%s
    \\multicolumn{2}{|r|}{\\textbf{Total:}} & \\textbf{%s} \\\\
    \\hline
\\end{tabular}
\\section*{Observações}
OBSERVACAO SERA ESCRITA EM BREVE
\\newpage
"""

latex_template_details = """
\\documentclass[a4paper,12pt]{article}
\\usepackage{fontspec}
\\setmainfont{Arial}
\\usepackage{polyglossia}
\\setmainlanguage{portuguese}
\\usepackage{geometry}
\\geometry{margin=2.2cm, headsep=1.5cm}
\\usepackage{graphicx}
\\usepackage{xcolor}
\\usepackage{colortbl}
\\usepackage{fancyhdr}
\\usepackage{parskip}
\\usepackage{lastpage}
\\usepackage{array}
\\usepackage{booktabs}
\\definecolor{readerred}{RGB}{207,45,75}
\\setlength{\\headheight}{70pt}
\\addtolength{\\topmargin}{-10pt}
\\pagestyle{fancy}
\\fancyhf{}
\\fancyhead[L]{\\color{readerred}\\small\\textbf{Insert Company Name} \\\\[-0.1cm] CNPJ: 12.345.567/0000-01 \\\\[-0.1cm] Insert Company Address Zip Code: 000-00 \\\\[-0.1cm] Company Phone Contact | Company Email Contact}
\\fancyhead[R]{\\raisebox{-0.5cm}{\\includegraphics[width=0.3\\textwidth]{logo.png}}}
\\fancyfoot[C]{\\thepage\\ de \\pageref{LastPage}}
\\newcolumntype{L}[1]{>{\\raggedright\\arraybackslash}p{#1}}
\\begin{document}
%s
\\end{document}
"""

latex_template_details_page = """
\\vspace*{-0.25cm}
\\begin{center}
    {\\LARGE \\textbf{Detalhes do Correspondente em %s}}\\\\
    \\vspace{0.5cm}
    {\\large Data: %s}
\\end{center}
\\section*{Informações do Correspondente}
\\begin{tabular}{|L{5cm}|L{9cm}|}
    \\hline
    \\rowcolor{readerred}
    \\color{white}\\textbf{Campo} & \\color{white}\\textbf{Valor} \\\\
    \\hline
    País de Atuação & %s \\\\
    Nome do Correspondente & %s \\\\
    Última Atualização & %s \\\\
    Peculiaridades de Trâmites & %s \\\\
    Descrição dos Documentos Necessários & %s \\\\
    \\hline
\\end{tabular}
\\newpage
"""

# Load Excel data
sheet_df = pd.read_excel(FILE_PATH)
sheet_df.columns = sheet_df.columns.str.strip()
sheet_df.columns = [unicodedata.normalize('NFKD', col).encode('ASCII', 'ignore').decode('ASCII').replace(" ", "_") for col in sheet_df.columns]
for col in sheet_df.columns:
    if col not in ['Ultima_Atualizacao', 'Nome_do_Correspondente', 'Pais_de_Atuacao', 'Peculiaridades_de_Tramites', 'Descricao_dos_Documentos_Necessarios']:
        sheet_df[col] = pd.to_numeric(sheet_df[col], errors='coerce')

# Date formatting
brasilia_tz = pytz.timezone('America/Sao_Paulo')
current_date = datetime.now(brasilia_tz)
day = f"{current_date.day:02d}"
month = month_names[current_date.month]
year = str(current_date.year)
formatted_date = f"{day} de {month} de {year}"
filename_date = current_date.strftime("%d-%m-%Y_%H-%M-%S")

# Setup logging
log_file = os.path.join(logs_dir, f"log_{filename_date}.txt")
logging.basicConfig(filename=log_file, level=logging.INFO, format='%(asctime)s - %(message)s')

def log_action(action):
    logging.info(action)

# Utility functions
def normalize_string(s):
    return unicodedata.normalize('NFKD', s).encode('ASCII', 'ignore').decode('ASCII').lower()

def escape_latex(text):
    if not isinstance(text, str):
        return text
    replacements = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_',
        '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}', '^': r'\textasciicircum{}', '\\': r'\textbackslash{}'
    }
    for char, escape in replacements.items():
        text = text.replace(char, escape)
    return text

def format_brazilian(value):
    if pd.isna(value):
        return "0,00"
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def get_trademark_type(stored_labels):
    for label in stored_labels:
        if 'Deposito_de_Marca_Mista_1_Classe' in label:
            return 'Mista'
        elif 'Deposito_de_Marca_Nominativa_1_Classe' in label:
            return 'Nominativa'
        elif 'Deposito_de_Marca_Figurativa_1_Classe' in label:
            return 'Figurativa'
    return ''


# Backend functions (unchanged from shell script)
def generate_marca_latex(marca_details, company_name):
    global generated_docx_files
    selected_country = marca_details['selected_country']
    selected_correspondent = marca_details['selected_correspondent']
    trademark_name = marca_details['trademark_name']
    deposit_option = marca_details['deposit_option']
    stored_labels = marca_details['stored_labels']
    additional_class_labels = marca_details['additional_class_labels']
    num_classes = marca_details['num_classes']

    row = sheet_df[(sheet_df['Pais_de_Atuacao'] == selected_country) & (sheet_df['Nome_do_Correspondente'] == selected_correspondent)].iloc[0]

    value_breakdown = {}
    latex_rows = []
    total_honorarios = 0
    total_deposito = 0
    total_concessao = 0
    total_prorrogacao = 0
    total_converted_sum = 0
    processed_labels = set()

    busca_labels = ['Busca'] if deposit_option == "Depósito" else []
    document_labels = ['Valor_Associados_a_Documentos_de_Deposito_de_Marca'] if deposit_option == "Depósito" else []
    deposito_labels = [label for label in stored_labels if "Deposito_de_Marca" in label]
    concessao_labels = [label for label in stored_labels if "Concessao_de_Marca" in label]
    prorrogacao_labels = [label for label in stored_labels if "Prorrogacao" in label]
    additional_deposito = [label for label in additional_class_labels if "Deposito_de_Marca" in label]
    additional_concessao = [label for label in additional_class_labels if "Concessao_de_Marca" in label]
    additional_prorrogacao = [label for label in additional_class_labels if "Prorrogacao" in label]

    def process_label(col, is_additional=False, multiplier=1):
        nonlocal total_honorarios, total_converted_sum, total_deposito, total_concessao, total_prorrogacao
        if col in row and pd.notnull(row[col]) and col not in processed_labels:
            extracted_val = row[col] * multiplier if pd.notnull(row[col]) else 0
            honor_key = f"Honorarios_{col}"
            honor_val = honorarios_dict.get(honor_key, 0) * multiplier
            if extracted_val != 0:
                total_val = (extracted_val + honor_val) * (1.235 if col not in ['Valor_Associados_a_Documentos_de_Deposito_de_Marca'] else 1)
            else:
                total_val = honor_val
            total_converted = total_val
            value_breakdown[col] = {
                'extracted_val': extracted_val,
                'honor_val': honor_val,
                'total_val': total_val,
                'total_converted': total_converted
            }
            total_honorarios += total_val
            total_converted_sum += total_converted
            if deposit_option == "Depósito" and "Concessao" not in col and col not in busca_labels:
                total_deposito += total_converted
            elif deposit_option == "Depósito" and "Concessao" in col:
                total_concessao += total_converted
            else:
                total_prorrogacao += total_converted
            description = col.replace('_', ' ').replace("Deposito", "Depósito").replace("Concessao", "Concessão").replace("Prorrogacao", "Prorrogação")
            valor_original = format_brazilian(extracted_val)
            total_converted_str = format_brazilian(total_converted)
            prefix = "\\textbf" if col in busca_labels or not is_additional else ""
            latex_rows.append(f"    {prefix}{{{description}}} & {valor_original} & {total_converted_str} \\\\ \\hline")
            processed_labels.add(col)

    for col in busca_labels:
        process_label(col)
    for col in document_labels:
        process_label(col)
    for col in deposito_labels:
        process_label(col)
    if num_classes > 1:
        additional_classes = num_classes - 1
        for col in additional_deposito:
            process_label(col, is_additional=True, multiplier=additional_classes)
    if deposit_option == "Depósito":
        value_breakdown["TOTAL_PARA_DEPOSITO"] = {'total_converted': total_deposito}
        latex_rows.append(f"    \\textbf{{TOTAL PARA DEPÓSITO DE MARCA}} & & \\textbf{{{format_brazilian(total_deposito)}}} \\\\ \\hline")
    for col in concessao_labels:
        process_label(col)
    if deposit_option == "Depósito" and num_classes > 1:
        additional_classes = num_classes - 1
        for col in additional_concessao:
            process_label(col, is_additional=True, multiplier=additional_classes)
    if deposit_option == "Depósito":
        value_breakdown["TOTAL_ESTIMADO_PARA_CONCESSAO"] = {'total_converted': total_concessao}
        latex_rows.append(f"    \\textbf{{TOTAL ESTIMADO PARA CONCESSÃO}} & & \\textbf{{{format_brazilian(total_concessao)}}} \\\\ \\hline")
    for col in prorrogacao_labels:
        process_label(col)
    if num_classes > 1:
        additional_classes = num_classes - 1
        for col in additional_prorrogacao:
            process_label(col, is_additional=True, multiplier=additional_classes)
    if deposit_option != "Depósito":
        value_breakdown["TOTAL_PARA_PRORROGACAO"] = {'total_converted': total_prorrogacao}
        latex_rows.append(f"    \\textbf{{TOTAL PARA PRORROGAÇÃO}} & & \\textbf{{{format_brazilian(total_prorrogacao)}}} \\\\ \\hline")

    latex_table = "\n".join(latex_rows)
    total_converted_str = format_brazilian(total_converted_sum)
    latex_page_template = latex_template_deposito_page if deposit_option == "Depósito" else latex_template_prorrogacao_page
    budget_latex_page = latex_page_template % (selected_country, formatted_date, company_name, trademark_name, latex_table, total_converted_str)

    details = {
        'pais_de_atuacao': escape_latex(row['Pais_de_Atuacao']),
        'nome_do_correspondente': escape_latex(row['Nome_do_Correspondente']),
        'ultima_atualizacao': escape_latex(row['Ultima_Atualizacao']),
        'peculiaridades_de_tramites': escape_latex(row['Peculiaridades_de_Tramites']),
        'descricao_dos_documentos_necessarios': escape_latex(row['Descricao_dos_Documentos_Necessarios'])
    }
    details_latex_page = latex_template_details_page % (
        selected_country, formatted_date, details['pais_de_atuacao'], details['nome_do_correspondente'],
        details['ultima_atualizacao'], details['peculiaridades_de_tramites'], details['descricao_dos_documentos_necessarios']
    )

    return budget_latex_page, details_latex_page, {
        'marca_details': marca_details,
        'company_name': company_name,
        'total_converted_sum': total_converted_sum,
        'value_breakdown': value_breakdown,
        'stored_labels': stored_labels,
        'total_deposito': total_deposito,
        'total_concessao': total_concessao
    }

def generate_docx_from_template(marca_details, company_name, total_converted_sum, value_breakdown, stored_labels, total_deposito, total_concessao):
    doc = Document(DOCX_TEMPLATE_PATH)
    def get_value(service, key, default=0):
        return value_breakdown.get(service, {}).get(key, default)
    trademark_type = get_trademark_type(stored_labels).lower()
    deposito_keyword = "Mista" if "mista" in trademark_type else "Nominativa" if "nominativa" in trademark_type else ""
    deposito_label = next(
        (label for label in stored_labels if f"Deposito_de_Marca_{deposito_keyword}_1_Classe" in label),
        next((label for label in stored_labels if "Deposito_de_Marca" in label), None)
    )
    for paragraph in doc.paragraphs:
        if '{{ pais }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
        if '{{ data }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
        if '{{ titular }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
        if '{{ marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
        if '{{ taxa_busca_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ taxa_busca_usd }}', format_brazilian(get_value('Busca', 'extracted_val')))
        if '{{ total_busca_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_busca_usd }}', format_brazilian(get_value('Busca', 'total_converted')))
        if '{{ taxa_documentos_usd }}' in paragraph.text:
            document_fee = get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'extracted_val')
            paragraph.text = paragraph.text.replace('{{ taxa_documentos_usd }}', format_brazilian(document_fee)).replace('{{ taxa_documentos_usd }} ', '')
        if '{{ total_documentos_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_documentos_usd }}', format_brazilian(get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'total_converted')))
        if '{{ tipo_marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ tipo_marca }}', get_trademark_type(stored_labels))
        if '{{ taxa_deposito_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ taxa_deposito_usd }}', format_brazilian(get_value(deposito_label, 'extracted_val')))
        if '{{ total_deposito_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_deposito_usd }}', format_brazilian(get_value(deposito_label, 'total_converted')))
        if '{{ total_para_deposito_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_para_deposito_usd }}', format_brazilian(total_deposito))
        if '{{ taxa_concessao_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ taxa_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'extracted_val')))
        if '{{ total_concessao_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'total_converted')))
        if '{{ total_estimado_concessao_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_estimado_concessao_usd }}', format_brazilian(total_concessao))
        if '{{ total_final_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_final_usd }}', format_brazilian(total_converted_sum))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{ pais }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
                    if '{{ data }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
                    if '{{ titular }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
                    if '{{ marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
                    if '{{ taxa_busca_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ taxa_busca_usd }}', format_brazilian(get_value('Busca', 'extracted_val')))
                    if '{{ total_busca_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_busca_usd }}', format_brazilian(get_value('Busca', 'total_converted')))
                    if '{{ taxa_documentos_usd }}' in paragraph.text:
                        document_fee = get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'extracted_val')
                        paragraph.text = paragraph.text.replace('{{ taxa_documentos_usd }}', format_brazilian(document_fee)).replace('{{ taxa_documentos_usd }} ', '')
                    if '{{ total_documentos_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_documentos_usd }}', format_brazilian(get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'total_converted')))
                    if '{{ tipo_marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ tipo_marca }}', get_trademark_type(stored_labels))
                    if '{{ taxa_deposito_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ taxa_deposito_usd }}', format_brazilian(get_value(deposito_label, 'extracted_val')))
                    if '{{ total_deposito_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_deposito_usd }}', format_brazilian(get_value(deposito_label, 'total_converted')))
                    if '{{ total_para_deposito_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_para_deposito_usd }}', format_brazilian(total_deposito))
                    if '{{ taxa_concessao_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ taxa_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'extracted_val')))
                    if '{{ total_concessao_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'total_converted')))
                    if '{{ total_estimado_concessao_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_estimado_concessao_usd }}', format_brazilian(total_concessao))
                    if '{{ total_final_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_final_usd }}', format_brazilian(total_converted_sum))
    safe_trademark_name = re.sub(r'[^\w\-]', '', marca_details['trademark_name'].replace(' ', ''))
    safe_country = re.sub(r'[^\w\-]', '', marca_details['selected_country'].replace(' ', ''))
    trademark_type = get_trademark_type(stored_labels)
    docx_filename = f"Orçamento_{safe_trademark_name}_{trademark_type}_{safe_country}_{filename_date}.docx"
    docx_output_path = os.path.join(output_dir, docx_filename)
    doc.save(docx_output_path)
    print(f"DOCX file successfully generated and saved to: {docx_output_path}")
    return docx_filename

def generate_docx_from_multi_class_template(marca_details, company_name, total_converted_sum, value_breakdown, stored_labels, total_deposito, total_concessao):
    doc = Document(DOCX_MULTI_CLASS_TEMPLATE_PATH)
    def get_value(service, key, default=0):
        return value_breakdown.get(service, {}).get(key, default)
    trademark_type = get_trademark_type(stored_labels).lower()
    deposito_keyword = "Mista" if "mista" in trademark_type else "Nominativa" if "nominativa" in trademark_type else ""
    deposito_label = next(
        (label for label in stored_labels if f"Deposito_de_Marca_{deposito_keyword}_1_Classe" in label),
        next((label for label in stored_labels if "Deposito_de_Marca" in label), None)
    )
    for paragraph in doc.paragraphs:
        if '{{ pais }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
        if '{{ data }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
        if '{{ titular }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
        if '{{ marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
        if '{{ taxa_busca_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ taxa_busca_usd }}', format_brazilian(get_value('Busca', 'extracted_val')))
        if '{{ total_busca_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_busca_usd }}', format_brazilian(get_value('Busca', 'total_converted')))
        if '{{ taxa_documentos_usd }}' in paragraph.text:
            document_fee = get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'extracted_val')
            paragraph.text = paragraph.text.replace('{{ taxa_documentos_usd }}', format_brazilian(document_fee)).replace('{{ taxa_documentos_usd }} ', '')
        if '{{ total_documentos_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_documentos_usd }}', format_brazilian(get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'total_converted')))
        if '{{ tipo_marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ tipo_marca }}', get_trademark_type(stored_labels))
        if '{{ taxa_deposito_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ taxa_deposito_usd }}', format_brazilian(get_value(deposito_label, 'extracted_val')))
        if '{{ total_deposito_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_deposito_usd }}', format_brazilian(get_value(deposito_label, 'total_converted')))
        if '{{ total_para_deposito_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_para_deposito_usd }}', format_brazilian(total_deposito))
        if '{{ taxa_concessao_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ taxa_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'extracted_val')))
        if '{{ total_concessao_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'total_converted')))
        if '{{ total_estimado_concessao_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_estimado_concessao_usd }}', format_brazilian(total_concessao))
        if '{{ total_final_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_final_usd }}', format_brazilian(total_converted_sum))
        if '{{ taxa_deposito_classe_adicional_usd }}' in paragraph.text:
            additional_deposito_label = next((label for label in marca_details['additional_class_labels'] if "Deposito_de_Marca" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_deposito_classe_adicional_usd }}', format_brazilian(get_value(additional_deposito_label, 'extracted_val')))
        if '{{ total_deposito_classe_adicional_usd }}' in paragraph.text:
            additional_deposito_label = next((label for label in marca_details['additional_class_labels'] if "Deposito_de_Marca" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_deposito_classe_adicional_usd }}', format_brazilian(get_value(additional_deposito_label, 'total_converted')))
        if '{{ taxa_concessao_classe_adicional_usd }}' in paragraph.text:
            additional_concessao_label = next((label for label in marca_details['additional_class_labels'] if "Concessao_de_Marca" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_concessao_classe_adicional_usd }}', format_brazilian(get_value(additional_concessao_label, 'extracted_val')))
        if '{{ total_concessao_classe_adicional_usd }}' in paragraph.text:
            additional_concessao_label = next((label for label in marca_details['additional_class_labels'] if "Concessao_de_Marca" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_concessao_classe_adicional_usd }}', format_brazilian(get_value(additional_concessao_label, 'total_converted')))
        if '{{ total_estimado_c_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_estimado_c_usd }}', format_brazilian(total_concessao))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{ pais }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
                    if '{{ data }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
                    if '{{ titular }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
                    if '{{ marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
                    if '{{ taxa_busca_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ taxa_busca_usd }}', format_brazilian(get_value('Busca', 'extracted_val')))
                    if '{{ total_busca_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_busca_usd }}', format_brazilian(get_value('Busca', 'total_converted')))
                    if '{{ taxa_documentos_usd }}' in paragraph.text:
                        document_fee = get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'extracted_val')
                        paragraph.text = paragraph.text.replace('{{ taxa_documentos_usd }}', format_brazilian(document_fee)).replace('{{ taxa_documentos_usd }} ', '')
                    if '{{ total_documentos_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_documentos_usd }}', format_brazilian(get_value('Valor_Associados_a_Documentos_de_Deposito_de_Marca', 'total_converted')))
                    if '{{ tipo_marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ tipo_marca }}', get_trademark_type(stored_labels))
                    if '{{ taxa_deposito_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ taxa_deposito_usd }}', format_brazilian(get_value(deposito_label, 'extracted_val')))
                    if '{{ total_deposito_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_deposito_usd }}', format_brazilian(get_value(deposito_label, 'total_converted')))
                    if '{{ total_para_deposito_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_para_deposito_usd }}', format_brazilian(total_deposito))
                    if '{{ taxa_concessao_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ taxa_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'extracted_val')))
                    if '{{ total_concessao_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_concessao_usd }}', format_brazilian(get_value('Concessao_de_Marca_1_Classe', 'total_converted')))
                    if '{{ total_estimado_concessao_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_estimado_concessao_usd }}', format_brazilian(total_concessao))
                    if '{{ total_final_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_final_usd }}', format_brazilian(total_converted_sum))
                    if '{{ taxa_deposito_classe_adicional_usd }}' in paragraph.text:
                        additional_deposito_label = next((label for label in marca_details['additional_class_labels'] if "Deposito_de_Marca" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_deposito_classe_adicional_usd }}', format_brazilian(get_value(additional_deposito_label, 'extracted_val')))
                    if '{{ total_deposito_classe_adicional_usd }}' in paragraph.text:
                        additional_deposito_label = next((label for label in marca_details['additional_class_labels'] if "Deposito_de_Marca" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_deposito_classe_adicional_usd }}', format_brazilian(get_value(additional_deposito_label, 'total_converted')))
                    if '{{ taxa_concessao_classe_adicional_usd }}' in paragraph.text:
                        additional_concessao_label = next((label for label in marca_details['additional_class_labels'] if "Concessao_de_Marca" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_concessao_classe_adicional_usd }}', format_brazilian(get_value(additional_concessao_label, 'extracted_val')))
                    if '{{ total_concessao_classe_adicional_usd }}' in paragraph.text:
                        additional_concessao_label = next((label for label in marca_details['additional_class_labels'] if "Concessao_de_Marca" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_concessao_classe_adicional_usd }}', format_brazilian(get_value(additional_concessao_label, 'total_converted')))
                    if '{{ total_estimado_c_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_estimado_c_usd }}', format_brazilian(total_concessao))
    safe_trademark_name = re.sub(r'[^\w\-]', '', marca_details['trademark_name'].replace(' ', ''))
    safe_country = re.sub(r'[^\w\-]', '', marca_details['selected_country'].replace(' ', ''))
    trademark_type = get_trademark_type(stored_labels)
    docx_filename = f"Orçamento_{safe_trademark_name}_{trademark_type}_{safe_country}_{filename_date}.docx"
    docx_output_path = os.path.join(output_dir, docx_filename)
    doc.save(docx_output_path)
    print(f"DOCX file successfully generated and saved to: {docx_output_path}")
    return docx_filename

def generate_docx_from_prorrogacao_ordinario_template(marca_details, company_name, total_converted_sum, value_breakdown):
    doc = Document(DOCX_PRORROGACAO_ORDINARIO_TEMPLATE_PATH)
    def get_value(service, key, default=0):
        return value_breakdown.get(service, {}).get(key, default)
    for paragraph in doc.paragraphs:
        if '{{ pais }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
        if '{{ data }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
        if '{{ titular }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
        if '{{ marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
        if '{{ taxa_prorrogacao_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
        if '{{ total_prorrogacao_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
        if '{{ total_para_prorrogacao_ordinario_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_ordinario_usd }}', format_brazilian(total_converted_sum))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{ pais }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
                    if '{{ data }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
                    if '{{ titular }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
                    if '{{ marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
                    if '{{ taxa_prorrogacao_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
                    if '{{ total_prorrogacao_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
                    if '{{ total_para_prorrogacao_ordinario_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_ordinario_usd }}', format_brazilian(total_converted_sum))
    safe_trademark_name = re.sub(r'[^\w\-]', '', marca_details['trademark_name'].replace(' ', ''))
    safe_country = re.sub(r'[^\w\-]', '', marca_details['selected_country'].replace(' ', ''))
    docx_filename = f"Orçamento_{safe_trademark_name}_Prorrogacao_Ordinario_{safe_country}_{filename_date}.docx"
    docx_output_path = os.path.join(output_dir, docx_filename)
    doc.save(docx_output_path)
    print(f"DOCX file successfully generated and saved to: {docx_output_path}")
    return docx_filename

def generate_docx_from_prorrogacao_ordinario_multi_class_template(marca_details, company_name, total_converted_sum, value_breakdown):
    doc = Document(DOCX_PRORROGACAO_ORDINARIO_MULTI_CLASS_TEMPLATE_PATH)
    def get_value(service, key, default=0):
        return value_breakdown.get(service, {}).get(key, default)
    for paragraph in doc.paragraphs:
        if '{{ pais }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
        if '{{ data }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
        if '{{ titular }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
        if '{{ marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
        if '{{ taxa_prorrogacao_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
        if '{{ total_prorrogacao_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
        if '{{ taxa_prorrogacao_classe_adicional_ordinario_usd }}' in paragraph.text:
            additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_classe_adicional_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'extracted_val')))
        if '{{ total_prorrogacao_classe_adicional_ordinario_usd }}' in paragraph.text:
            additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_prorrogacao_classe_adicional_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'total_converted')))
        if '{{ total_para_prorrogacao_ordinario_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_ordinario_usd }}', format_brazilian(total_converted_sum))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{ pais }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
                    if '{{ data }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
                    if '{{ titular }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
                    if '{{ marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
                    if '{{ taxa_prorrogacao_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
                    if '{{ total_prorrogacao_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_prorrogacao_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
                    if '{{ taxa_prorrogacao_classe_adicional_ordinario_usd }}' in paragraph.text:
                        additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_classe_adicional_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'extracted_val')))
                    if '{{ total_prorrogacao_classe_adicional_ordinario_usd }}' in paragraph.text:
                        additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_prorrogacao_classe_adicional_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'total_converted')))
                    if '{{ total_para_prorrogacao_ordinario_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_ordinario_usd }}', format_brazilian(total_converted_sum))
    safe_trademark_name = re.sub(r'[^\w\-]', '', marca_details['trademark_name'].replace(' ', ''))
    safe_country = re.sub(r'[^\w\-]', '', marca_details['selected_country'].replace(' ', ''))
    docx_filename = f"Orçamento_{safe_trademark_name}_Prorrogacao_Ordinario_{safe_country}_{filename_date}.docx"
    docx_output_path = os.path.join(output_dir, docx_filename)
    doc.save(docx_output_path)
    print(f"DOCX file successfully generated and saved to: {docx_output_path}")
    return docx_filename

def generate_docx_from_prorrogacao_extra_ordinario_template(marca_details, company_name, total_converted_sum, value_breakdown):
    doc = Document(DOCX_PRORROGACAO_EXTRA_ORDINARIO_TEMPLATE_PATH)
    def get_value(service, key, default=0):
        return value_breakdown.get(service, {}).get(key, default)
    for paragraph in doc.paragraphs:
        if '{{ pais }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
        if '{{ data }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
        if '{{ titular }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
        if '{{ marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
        if '{{ taxa_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
        if '{{ total_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
        if '{{ total_para_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_extra_ordinario_usd }}', format_brazilian(total_converted_sum))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{ pais }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
                    if '{{ data }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
                    if '{{ titular }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
                    if '{{ marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
                    if '{{ taxa_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
                    if '{{ total_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
                    if '{{ total_para_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_extra_ordinario_usd }}', format_brazilian(total_converted_sum))
    safe_trademark_name = re.sub(r'[^\w\-]', '', marca_details['trademark_name'].replace(' ', ''))
    safe_country = re.sub(r'[^\w\-]', '', marca_details['selected_country'].replace(' ', ''))
    docx_filename = f"Orçamento_{safe_trademark_name}_Prorrogacao_Extra_Ordinario_{safe_country}_{filename_date}.docx"
    docx_output_path = os.path.join(output_dir, docx_filename)
    doc.save(docx_output_path)
    print(f"DOCX file successfully generated and saved to: {docx_output_path}")
    return docx_filename

def generate_docx_from_prorrogacao_extra_ordinario_multi_class_template(marca_details, company_name, total_converted_sum, value_breakdown):
    doc = Document(DOCX_PRORROGACAO_EXTRA_ORDINARIO_MULTI_CLASS_TEMPLATE_PATH)
    def get_value(service, key, default=0):
        return value_breakdown.get(service, {}).get(key, default)
    for paragraph in doc.paragraphs:
        if '{{ pais }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
        if '{{ data }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
        if '{{ titular }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
        if '{{ marca }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
        if '{{ taxa_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
        if '{{ total_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
            prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
        if '{{ taxa_prorrogacao_classe_adicional_extra_ordinario_usd }}' in paragraph.text:
            additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_classe_adicional_extra_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'extracted_val')))
        if '{{ total_prorrogacao_classe_adicional_extra_ordinario_usd }}' in paragraph.text:
            additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
            paragraph.text = paragraph.text.replace('{{ total_prorrogacao_classe_adicional_extra_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'total_converted')))
        if '{{ total_para_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_extra_ordinario_usd }}', format_brazilian(total_converted_sum))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{ pais }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ pais }}', marca_details['selected_country'])
                    if '{{ data }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ data }}', formatted_date)
                    if '{{ titular }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ titular }}', company_name)
                    if '{{ marca }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ marca }}', marca_details['trademark_name'])
                    if '{{ taxa_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'extracted_val')))
                    if '{{ total_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
                        prorrogacao_label = next((label for label in marca_details['stored_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_prorrogacao_extra_ordinario_usd }}', format_brazilian(get_value(prorrogacao_label, 'total_converted')))
                    if '{{ taxa_prorrogacao_classe_adicional_extra_ordinario_usd }}' in paragraph.text:
                        additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ taxa_prorrogacao_classe_adicional_extra_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'extracted_val')))
                    if '{{ total_prorrogacao_classe_adicional_extra_ordinario_usd }}' in paragraph.text:
                        additional_prorrogacao_label = next((label for label in marca_details['additional_class_labels'] if "Prorrogacao_Prazo_Extra_Ordinario" in label), None)
                        paragraph.text = paragraph.text.replace('{{ total_prorrogacao_classe_adicional_extra_ordinario_usd }}', format_brazilian(get_value(additional_prorrogacao_label, 'total_converted')))
                    if '{{ total_para_prorrogacao_extra_ordinario_usd }}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{ total_para_prorrogacao_extra_ordinario_usd }}', format_brazilian(total_converted_sum))
    safe_trademark_name = re.sub(r'[^\w\-]', '', marca_details['trademark_name'].replace(' ', ''))
    safe_country = re.sub(r'[^\w\-]', '', marca_details['selected_country'].replace(' ', ''))
    docx_filename = f"Orçamento_{safe_trademark_name}_Prorrogacao_Extra_Ordinario_{safe_country}_{filename_date}.docx"
    docx_output_path = os.path.join(output_dir, docx_filename)
    doc.save(docx_output_path)
    print(f"DOCX file successfully generated and saved to: {docx_output_path}")
    return docx_filename

# Tkinter Application
class Application:
    def __init__(self, root):
        self.root = root
        self.root.title("M-G-L International Ferramenta de Orçamento de Marcas")
        self.root.geometry("900x600")
        self.root.configure(bg="#f1e7d4")

        self.inputs = []
        self.current_marca = {}
        self.navigation_stack = []
        self.generated_docx_files = []
        self.company_name = None
        self.same_country = None
        self.selected_country = None
        self.selected_correspondent = None
        self.global_country = None
        self.global_correspondent = None

        self.style = ttk.Style()
        self.style.configure("TButton", font=("Helvetica", 12), padding=12, foreground="#3e2212", background="#3e2212")
        self.style.configure("TLabel", font=("Helvetica", 14), background="#f1e7d4", foreground="#3e2212")
        self.style.configure("TEntry", font=("Helvetica", 12), padding=5, foreground="#3e2212")

        self.validate_positive_int = (self.root.register(self.validate_input), '%P')
        self.countries = sorted(sheet_df['Pais_de_Atuacao'].dropna().unique())
        self.PAISES = [f"{i+1}. {country}" for i, country in enumerate(self.countries)]

        log_action("Application initialized")
        # CHANGE 1: Start with the welcome screen instead of company name screen
        self.show_start_screen()

    def clear_frame(self):
        for widget in self.root.winfo_children():
            widget.destroy()

    def add_back_button(self, back_command):
        back_button = ttk.Button(self.root, text="Voltar", command=back_command)
        back_button.place(x=750, y=540)

    def validate_input(self, value):
        if value == "":
            return True
        try:
            num = int(value)
            return num > 0
        except ValueError:
            return False

    def show_start_screen(self):
        self.clear_frame()
        frame = tk.Frame(self.root, bg="#f1e7d4")
        frame.place(relx=0.5, rely=0.5, anchor="center")

        label = ttk.Label(frame, text="Bem-vindo à ferramenta de orçamento de marcas")
        label.pack(pady=40)

        # CHANGE 1: "COMEÇAR" now leads to company name input
        start_button = ttk.Button(frame, text="COMEÇAR", command=self.show_company_name_screen)
        start_button.pack(pady=20)

        exit_button = ttk.Button(frame, text="SAIR", command=self.root.quit)
        exit_button.pack(pady=10)
        log_action("Start screen displayed")

    def show_company_name_screen(self):
        # CHANGE 1: Navigation stack updated to reflect new flow (back to start screen)
        self.navigation_stack.append(self.show_start_screen)
        self.clear_frame()
        frame = tk.Frame(self.root, bg="#f1e7d4")
        frame.place(relx=0.5, rely=0.5, anchor="center")

        label = ttk.Label(frame, text="Digite o Nome da Empresa (Titular)")
        label.pack(pady=20)

        self.company_entry = ttk.Entry(frame, width=50)
        self.company_entry.pack(pady=10)
        # CHANGE 2: Bind Enter key to save_company_name
        self.company_entry.bind('<Return>', lambda event: self.save_company_name())

        next_button = ttk.Button(frame, text="Avançar", command=self.save_company_name)
        next_button.pack(pady=20)

        self.add_back_button(self.navigation_stack[-1])

    def save_company_name(self):
        company_name = self.company_entry.get().strip()
        if company_name:
            self.company_name = escape_latex(company_name)
            log_action(f"Company name entered: {self.company_name}")
            self.show_multiple_brands_question()
        else:
            messagebox.showerror("Erro", "Por favor, insira o nome da empresa.")

    def show_multiple_brands_question(self):
        self.navigation_stack.append(self.show_company_name_screen)
        self.clear_frame()

        label = ttk.Label(self.root, text="Deseja Orçar Múltiplas Marcas?")
        label.place(relx=0.5, rely=0.3, anchor="center")

        yes_button = ttk.Button(self.root, text="Sim", command=lambda: self.save_multibrand_choice("Sim"))
        no_button = ttk.Button(self.root, text="Não", command=lambda: self.save_multibrand_choice("Não"))
        yes_button.place(relx=0.4, rely=0.4, anchor="center")
        no_button.place(relx=0.6, rely=0.4, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Multiple brands question displayed")

    def save_multibrand_choice(self, choice):
        self.current_marca["multiplas_marcas"] = choice
        log_action(f"Multiple brands choice: {choice}")
        if choice == "Sim":
            self.show_same_country_question()
        else:
            self.show_country_selection()

    def show_same_country_question(self):
        self.navigation_stack.append(self.show_multiple_brands_question)
        self.clear_frame()

        label = ttk.Label(self.root, text="Todas as Marcas São para o Mesmo País?")
        label.place(relx=0.5, rely=0.3, anchor="center")

        yes_button = ttk.Button(self.root, text="Sim", command=lambda: self.save_same_country("Sim"))
        no_button = ttk.Button(self.root, text="Não", command=lambda: self.save_same_country("Não"))
        yes_button.place(relx=0.4, rely=0.4, anchor="center")
        no_button.place(relx=0.6, rely=0.4, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Same country question displayed")

    def save_same_country(self, choice):
        self.same_country = choice
        log_action(f"Same country choice: {choice}")
        self.show_country_selection()

    def show_country_selection(self):
        self.navigation_stack.append(self.show_same_country_question if self.current_marca.get("multiplas_marcas") == "Sim" else self.show_multiple_brands_question)
        self.clear_frame()

        label = ttk.Label(self.root, text="Países Disponíveis (Digite ou selecione e pressione Enter)")
        label.place(relx=0.5, rely=0.05, anchor="center")

        self.country_entry = ttk.Entry(self.root, width=50)
        self.country_entry.place(relx=0.5, rely=0.1, anchor="center")
        # Bind Enter key to save_country for entry field
        self.country_entry.bind('<Return>', lambda event: self.save_country())

        list_frame = tk.Frame(self.root, bg="#f1e7d4")
        list_frame.place(relx=0.5, rely=0.55, anchor="center", width=600, height=300)

        scrollbar = tk.Scrollbar(list_frame, orient="vertical")
        scrollbar.pack(side="right", fill="y")

        self.country_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, font=("Helvetica", 10), height=15, width=60)
        self.country_listbox.pack(side="left", fill="both", expand=True)

        for country in self.PAISES:
            self.country_listbox.insert(tk.END, country)

        scrollbar.config(command=self.country_listbox.yview)
        self.country_listbox.bind('<<ListboxSelect>>', self.on_country_select)
        # Bind Enter key to save_country for listbox
        self.country_listbox.bind('<Return>', lambda event: self.save_country())

        next_button = ttk.Button(self.root, text="Avançar", command=self.save_country)
        next_button.place(relx=0.5, rely=0.9, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Country selection screen displayed")


    def on_country_select(self, event):
        selection = self.country_listbox.curselection()
        if selection:
            self.country_entry.delete(0, tk.END)
            self.country_entry.insert(0, self.country_listbox.get(selection[0]))

    def save_country(self):
        country_input = self.country_entry.get().strip()
        # Handle number input (e.g., "1")
        if country_input.isdigit() and 1 <= int(country_input) <= len(self.countries):
            self.selected_country = self.countries[int(country_input) - 1]
        # Handle name input (e.g., "Brasil")
        elif any(normalize_string(country_input) == normalize_string(country) for country in self.countries):
            self.selected_country = next((country for country in self.countries if normalize_string(country_input) == normalize_string(country)), None)
            if self.selected_country is None:
                messagebox.showerror("Erro", "País inválido. Digite o número, nome do país ou 'número. Nome' (ex: '1. Brasil').")
                return
        # Handle "number. Name" input (e.g., "2. AFRICA DO SUL")
        elif re.match(r'^\d+\.\s[A-Za-z\s]+$', country_input):
            try:
                number, name = country_input.split('. ', 1)
                number = int(number)
                if 1 <= number <= len(self.countries) and normalize_string(name) == normalize_string(self.countries[number - 1]):
                    self.selected_country = self.countries[number - 1]
                else:
                    messagebox.showerror("Erro", "País inválido. Digite o número, nome do país ou 'número. Nome' (ex: '1. Brasil').")
                    return
            except (ValueError, IndexError):
                messagebox.showerror("Erro", "País inválido. Digite o número, nome do país ou 'número. Nome' (ex: '1. Brasil').")
                return
        else:
            messagebox.showerror("Erro", "País inválido. Digite o número, nome do país ou 'número. Nome' (ex: '1. Brasil').")
            return
        self.current_marca["selected_country"] = self.selected_country
        log_action(f"Country selected: {self.selected_country}")
        self.show_correspondent_input()

    def show_correspondent_input(self):
        self.navigation_stack.append(self.show_country_selection)
        self.clear_frame()

        label = ttk.Label(self.root, text="Correspondentes Disponíveis (Digite ou selecione e pressione Enter)")
        label.place(relx=0.5, rely=0.05, anchor="center")

        self.correspondent_entry = ttk.Entry(self.root, width=50)
        self.correspondent_entry.place(relx=0.5, rely=0.1, anchor="center")
        # Bind Enter key to save_correspondent for entry field
        self.correspondent_entry.bind('<Return>', lambda event: self.save_correspondent())

        correspondents = sorted(sheet_df[sheet_df['Pais_de_Atuacao'] == self.selected_country]['Nome_do_Correspondente'].dropna().unique())
        self.correspondents = [f"{i+1}. {corr}" for i, corr in enumerate(correspondents)]

        list_frame = tk.Frame(self.root, bg="#f1e7d4")
        list_frame.place(relx=0.5, rely=0.55, anchor="center", width=600, height=300)

        scrollbar = tk.Scrollbar(list_frame, orient="vertical")
        scrollbar.pack(side="right", fill="y")

        self.correspondent_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, font=("Helvetica", 10), height=15, width=60)
        self.correspondent_listbox.pack(side="left", fill="both", expand=True)

        for corr in self.correspondents:
            self.correspondent_listbox.insert(tk.END, corr)

        scrollbar.config(command=self.correspondent_listbox.yview)
        self.correspondent_listbox.bind('<<ListboxSelect>>', self.on_correspondent_select)
        # Bind Enter key to save_correspondent for listbox
        self.correspondent_listbox.bind('<Return>', lambda event: self.save_correspondent())

        next_button = ttk.Button(self.root, text="Avançar", command=self.save_correspondent)
        next_button.place(relx=0.5, rely=0.9, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Correspondent selection screen displayed")

    def on_correspondent_select(self, event):
        selection = self.correspondent_listbox.curselection()
        if selection:
            self.correspondent_entry.delete(0, tk.END)
            self.correspondent_entry.insert(0, self.correspondent_listbox.get(selection[0]))

    def save_correspondent(self):
        corr_input = self.correspondent_entry.get().strip()
        correspondents = sorted(sheet_df[sheet_df['Pais_de_Atuacao'] == self.selected_country]['Nome_do_Correspondente'].dropna().unique())
        # Handle number input (e.g., "1")
        if corr_input.isdigit() and 1 <= int(corr_input) <= len(correspondents):
            self.selected_correspondent = correspondents[int(corr_input) - 1]
        # Handle name input (e.g., "John Doe")
        elif any(normalize_string(corr_input) == normalize_string(corr) for corr in correspondents):
            self.selected_correspondent = next((corr for corr in correspondents if normalize_string(corr_input) == normalize_string(corr)), None)
            if self.selected_correspondent is None:
                messagebox.showerror("Erro", "Correspondente inválido. Digite o número, nome do correspondente ou 'número. Nome' (ex: '1. John Doe').")
                return
        # Handle "number. Name" input (e.g., "1. John Doe" or "1. José Silva")
        elif re.match(r'^\d+\.\s[\w\sÀ-ÿ\'-]+$', corr_input):
            try:
                number, name = corr_input.split('. ', 1)
                number = int(number)
                if 1 <= number <= len(correspondents) and normalize_string(name) == normalize_string(correspondents[number - 1]):
                    self.selected_correspondent = correspondents[number - 1]
                else:
                    messagebox.showerror("Erro", "Correspondente inválido. Verifique se o número e o nome correspondem à lista (ex: '1. John Doe').")
                    return
            except (ValueError, IndexError):
                messagebox.showerror("Erro", "Formato inválido. Use 'número. Nome' com um espaço após o ponto (ex: '1. John Doe').")
                return
        else:
            messagebox.showerror("Erro", "Correspondente inválido. Digite o número, nome do correspondente ou 'número. Nome' (ex: '1. John Doe').")
            return
        self.current_marca["selected_correspondent"] = self.selected_correspondent
        log_action(f"Correspondent selected: {self.selected_correspondent}")
        if self.same_country == "Sim":
            self.global_country = self.selected_country
            self.global_correspondent = self.selected_correspondent
        self.show_trademark_name_input()

    def show_trademark_name_input(self):
        self.navigation_stack.append(self.show_correspondent_input)
        self.clear_frame()

        label = ttk.Label(self.root, text="Digite o Nome da Marca")
        label.place(relx=0.5, rely=0.3, anchor="center")

        self.trademark_entry = ttk.Entry(self.root, width=50)
        self.trademark_entry.place(relx=0.5, rely=0.4, anchor="center")
        # CHANGE 2: Bind Enter key to save_trademark_name
        self.trademark_entry.bind('<Return>', lambda event: self.save_trademark_name())

        next_button = ttk.Button(self.root, text="Avançar", command=self.save_trademark_name)
        next_button.place(relx=0.5, rely=0.5, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Trademark name input screen displayed")

    def save_trademark_name(self):
        trademark_name = self.trademark_entry.get().strip()
        if trademark_name:
            self.current_marca["trademark_name"] = escape_latex(trademark_name)
            log_action(f"Trademark name entered: {trademark_name}")
            self.show_options()
        else:
            messagebox.showerror("Erro", "Por favor, insira o nome da marca.")

    def show_options(self):
        self.navigation_stack.append(self.show_trademark_name_input)
        self.clear_frame()

        label = ttk.Label(self.root, text="Selecione o Serviço a ser Orçado")
        label.place(relx=0.5, rely=0.2, anchor="center")

        ttk.Button(self.root, text="Depósito", command=lambda: self.set_service("Depósito")).place(relx=0.5, rely=0.35, anchor="center")
        ttk.Button(self.root, text="Prorrogação Ordinária", command=lambda: self.set_service("Prorrogação Ordinária")).place(relx=0.5, rely=0.45, anchor="center")
        ttk.Button(self.root, text="Prorrogação Extraordinária", command=lambda: self.set_service("Prorrogação Extraordinária")).place(relx=0.5, rely=0.55, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Service selection screen displayed")

    def set_service(self, service):
        if service == "Depósito":
            self.current_marca["deposit_option"] = "Depósito"
            log_action("Service selected: Depósito")
            self.show_trademark_type()
        else:
            self.current_marca["deposit_option"] = "Prorrogação"
            self.current_marca["prorrogacao_type"] = "Prazo_Ordinario" if service == "Prorrogação Ordinária" else "Prazo_Extra_Ordinario"
            self.current_marca["stored_labels"] = [f"Prorrogacao_{self.current_marca['prorrogacao_type']}_1_Classe"]
            self.current_marca["additional_class_labels"] = []
            log_action(f"Service selected: {service}")
            self.show_class_count()

    def show_trademark_type(self):
        self.navigation_stack.append(self.show_options)
        self.clear_frame()

        label = ttk.Label(self.root, text="Selecione o Tipo de Marca")
        label.place(relx=0.5, rely=0.2, anchor="center")

        ttk.Button(self.root, text="Nominativa", command=lambda: self.save_trademark_type("Nominativa")).place(relx=0.5, rely=0.35, anchor="center")
        ttk.Button(self.root, text="Figurativa", command=lambda: self.save_trademark_type("Figurativa")).place(relx=0.5, rely=0.45, anchor="center")
        ttk.Button(self.root, text="Mista", command=lambda: self.save_trademark_type("Mista")).place(relx=0.5, rely=0.55, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Trademark type selection screen displayed")

    def save_trademark_type(self, tipo):
        self.current_marca["tipo_marca"] = tipo
        self.current_marca["stored_labels"] = [
            "Busca",
            f"Deposito_de_Marca_{tipo}_1_Classe",
            "Concessao_de_Marca_1_Classe"
        ]
        self.current_marca["additional_class_labels"] = []
        log_action(f"Trademark type selected: {tipo}")
        self.show_class_count()

    def show_class_count(self):
        self.navigation_stack.append(self.show_trademark_type if self.current_marca["deposit_option"] == "Depósito" else self.show_options)
        self.clear_frame()

        label = ttk.Label(self.root, text="Digite o Número de Classes")
        label.place(relx=0.5, rely=0.3, anchor="center")

        self.classes_entry = ttk.Entry(self.root, width=20, validate="key", validatecommand=self.validate_positive_int)
        self.classes_entry.place(relx=0.5, rely=0.4, anchor="center")
        # CHANGE 2: Bind Enter key to save_class_count
        self.classes_entry.bind('<Return>', lambda event: self.save_class_count())

        next_button = ttk.Button(self.root, text="Avançar", command=self.save_class_count)
        next_button.place(relx=0.5, rely=0.5, anchor="center")

        self.add_back_button(self.navigation_stack[-1])
        log_action("Class count input screen displayed")

    def save_class_count(self):
        valor = self.classes_entry.get().strip()
        log_action(f"Class count input received: '{valor}'")
        if self.validate_input(valor) and valor:
            try:
                num_classes = int(valor)
                self.current_marca["num_classes"] = num_classes
                if num_classes > 1:
                    if self.current_marca["deposit_option"] == "Depósito":
                        self.current_marca["additional_class_labels"] = [
                            f"Deposito_de_Marca_{self.current_marca['tipo_marca']}_Classe_Adicional",
                            "Concessao_de_Marca_Classe_Adicional"
                        ]
                    else:
                        self.current_marca["additional_class_labels"] = [
                            f"Prorrogacao_{self.current_marca['prorrogacao_type']}_Classe_Adicional"
                        ]
                self.inputs.append(self.current_marca.copy())
                log_action(f"Class count validated and saved: {num_classes} for trademark {self.current_marca['trademark_name']}")
                self.check_for_more_trademarks()
            except ValueError as e:
                log_action(f"Error converting class count: {e}")
                messagebox.showerror("Erro", "Número de classes inválido. Insira um número positivo.")
        else:
            log_action("Invalid class count input")
            messagebox.showerror("Erro", "Por favor, insira um número positivo válido para classes.")

    def check_for_more_trademarks(self):
        log_action(f"Checking for more trademarks, multiplas_marcas: {self.current_marca.get('multiplas_marcas')}")
        if self.current_marca.get("multiplas_marcas") == "Sim":
            self.clear_frame()
            label = ttk.Label(self.root, text="Deseja Adicionar Outra Marca?")
            label.place(relx=0.5, rely=0.3, anchor="center")

            yes_button = ttk.Button(self.root, text="Sim", command=self.add_another_trademark)
            no_button = ttk.Button(self.root, text="Não", command=self.finalize_budget)
            yes_button.place(relx=0.4, rely=0.4, anchor="center")
            no_button.place(relx=0.6, rely=0.4, anchor="center")

            self.add_back_button(self.show_class_count)
            log_action("Prompt for additional trademark displayed")
        else:
            log_action("Proceeding to finalize budget (single trademark)")
            self.finalize_budget()

    def add_another_trademark(self):
        log_action("User chose to add another trademark")
        self.current_marca = {"multiplas_marcas": "Sim"}
        if self.same_country == "Sim":
            self.current_marca["selected_country"] = self.global_country
            self.current_marca["selected_correspondent"] = self.global_correspondent
            log_action("Reusing same country/correspondent")
            self.show_trademark_name_input()
        else:
            log_action("Selecting new country/correspondent")
            self.show_country_selection()

    def finalize_budget(self):
        log_action("Starting budget finalization")
        self.clear_frame()
        processing_label = ttk.Label(self.root, text="Processando arquivos, por favor aguarde...")
        processing_label.place(relx=0.5, rely=0.5, anchor="center")
        self.root.update()

        errors = []
        try:
            # Copy logo.png to logs_dir
            logo_output_path = os.path.join(logs_dir, "logo.png")
            if os.path.exists(LOGO_PATH):
                shutil.copy(LOGO_PATH, logo_output_path)
                log_action(f"Copied logo.png to {logo_output_path}")
            else:
                log_action(f"logo.png not found at {LOGO_PATH}")
                errors.append(f"Arquivo logo.png não encontrado em {LOGO_PATH}")

            latex_pages = []
            details_pages = []
            docx_data_list = []
            for marca_details in self.inputs:
                if "selected_country" not in marca_details:
                    marca_details["selected_country"] = self.global_country
                if "selected_correspondent" not in marca_details:
                    marca_details["selected_correspondent"] = self.global_correspondent
                log_action(f"Processing trademark: {marca_details.get('trademark_name')}")
                budget_page, details_page, docx_data = generate_marca_latex(marca_details, self.company_name)
                latex_pages.append(budget_page)
                details_pages.append(details_page)
                docx_data_list.append(docx_data)
                log_action(f"Generated LaTeX and DOCX data for trademark: {marca_details.get('trademark_name')}")

            # Generate LaTeX files in logs_dir
            full_latex_budget = latex_template_deposito % "\n".join(latex_pages)
            full_latex_details = latex_template_details % "\n".join(details_pages)
            budget_tex_filename = f"Orçamento_{filename_date}.tex"
            details_tex_filename = f"Detalhes_{filename_date}.tex"
            budget_tex_path = os.path.join(logs_dir, budget_tex_filename)
            details_tex_path = os.path.join(logs_dir, details_tex_filename)
            try:
                with open(budget_tex_path, "w", encoding="utf-8") as f:
                    f.write(full_latex_budget)
                with open(details_tex_path, "w", encoding="utf-8") as f:
                    f.write(full_latex_details)
                log_action(f"LaTeX files saved: {budget_tex_path}, {details_tex_path}")
                log_action(f"Budget LaTeX file size: {os.path.getsize(budget_tex_path)} bytes")
                log_action(f"Details LaTeX file size: {os.path.getsize(details_tex_path)} bytes")
                log_action(f"Budget LaTeX snippet: {full_latex_budget[:200]}...")
                log_action(f"Details LaTeX snippet: {full_latex_details[:200]}...")
            except IOError as e:
                log_action(f"Error saving LaTeX files: {e}")
                errors.append(f"Erro ao salvar arquivos LaTeX: {e}")

            # Generate DOCX files
            for docx_data in docx_data_list:
                log_action(f"Starting DOCX generation for trademark: {docx_data['marca_details'].get('trademark_name')}")
                try:
                    marca_details = docx_data['marca_details']
                    total_converted_sum = docx_data['total_converted_sum']
                    value_breakdown = docx_data['value_breakdown']
                    stored_labels = docx_data['stored_labels']
                    total_deposito = docx_data['total_deposito']
                    total_concessao = docx_data['total_concessao']
                    if marca_details['deposit_option'] == "Depósito":
                        if marca_details['num_classes'] > 1:
                            docx_filename = generate_docx_from_multi_class_template(
                                marca_details, self.company_name, total_converted_sum, value_breakdown, stored_labels, total_deposito, total_concessao
                            )
                        else:
                            docx_filename = generate_docx_from_template(
                                marca_details, self.company_name, total_converted_sum, value_breakdown, stored_labels, total_deposito, total_concessao
                            )
                    elif marca_details['prorrogacao_type'] == "Prazo_Ordinario":
                        if marca_details['num_classes'] > 1:
                            docx_filename = generate_docx_from_prorrogacao_ordinario_multi_class_template(
                                marca_details, self.company_name, total_converted_sum, value_breakdown
                            )
                        else:
                            docx_filename = generate_docx_from_prorrogacao_ordinario_template(
                                marca_details, self.company_name, total_converted_sum, value_breakdown
                            )
                    else:
                        if marca_details['num_classes'] > 1:
                            docx_filename = generate_docx_from_prorrogacao_extra_ordinario_multi_class_template(
                                marca_details, self.company_name, total_converted_sum, value_breakdown
                            )
                        else:
                            docx_filename = generate_docx_from_prorrogacao_extra_ordinario_template(
                                marca_details, self.company_name, total_converted_sum, value_breakdown
                            )
                    self.generated_docx_files.append(docx_filename)
                    log_action(f"DOCX generated: {docx_filename}")
                except Exception as e:
                    log_action(f"Error generating DOCX: {e}")
                    errors.append(f"Erro ao gerar arquivo DOCX: {e}")

            # Compile LaTeX to PDF
            if shutil.which("xelatex"):
                for tex_filename, output_name in [(budget_tex_filename, f"Orçamento_{filename_date}.pdf"), (details_tex_filename, f"Detalhes_{filename_date}.pdf")]:
                    log_action(f"Starting LaTeX compilation for {tex_filename}")
                    try:
                        result = subprocess.run(
                            ['xelatex', '-interaction=nonstopmode', tex_filename],
                            cwd=logs_dir,
                            capture_output=True,
                            encoding='utf-8'
                        )
                        log_action(f"LaTeX compilation output: {result.stdout[:200]}...")
                        if result.stderr:
                            log_action(f"LaTeX compilation warnings/errors: {result.stderr[:200]}...")
                        pdf_path = os.path.join(pdf_output_dir, output_name)
                        if os.path.exists(os.path.join(logs_dir, output_name)):
                            shutil.move(os.path.join(logs_dir, output_name), pdf_path)
                            log_action(f"PDF moved to: {pdf_path}")
                        else:
                            log_action(f"Failed to generate PDF: {output_name}")
                            errors.append(f"Falha ao gerar PDF: {output_name}")
                    except subprocess.CalledProcessError as e:
                        log_action(f"Error compiling LaTeX for {output_name}: {e.stderr[:200]}...")
                        errors.append(f"Erro processing PDF {output_name}: {e}")
                    except Exception as e:
                        log_action(f"Unexpected error during LaTeX compilation for {output_name}: {e}")
                        errors.append(f"Erro inesperado ao compilar PDF {output_name}: {e}")

            else:
                log_action("XeLaTeX not found, skipping PDF compilation")
                errors.append("XeLaTeX not found. PDFs não foram gerados.")

            if errors:
                messagebox.showerror("Erro", "\n".join(errors))
                self.show_completion_screen(error=True)
            else:
                log_action("Budget generation completed successfully")
                self.show_completion_screen(error=False)

        except Exception as e:
            log_action(f"Unexpected error in finalize_budget: {e}")
            messagebox.showerror("Erro", f"Erro inesperado: {e}")
            self.show_completion_screen(error=True)

    def show_completion_screen(self, error=False):
        self.clear_frame()
        frame = tk.Frame(self.root, bg="#f1e7d4")
        frame.place(relx=0.5, rely=0.5, anchor="center")
        if error:
            label = ttk.Label(frame, text="Erro ao gerar o orçamento. Verifique o log para detalhes.")
        else:
            label = ttk.Label(frame, text=f"Orçamento Gerado com Sucesso! Arquivos salvos em: {logs_dir}")
        label.pack(pady=20)
        new_budget_button = ttk.Button(frame, text="Novo Orçamento", command=self.reset_application)
        new_budget_button.pack(pady=10)
        exit_button = ttk.Button(frame, text="SAIR", command=self.root.quit)
        exit_button.pack(pady=10)
        label.pack(pady=10)
        log_action(f"Completion screen displayed, error: {error}")

    def reset_application(self):
        self.inputs = []
        self.current_marca = {}
        self.navigation_stack = []
        self.generated_docx_files = []
        self.company_name = None
        self.same_country = None
        self.selected_country = None
        self.selected_correspondent = None
        self.global_country = None
        self.global_correspondent = None
        log_action("Application reset for new budget")
        # CHANGE 1: Reset to start screen for new budget
        self.show_start_screen()

# Main application start
if __name__ == "__main__":
    root = tk.Tk()
    app = Application(root)
    root.mainloop()
